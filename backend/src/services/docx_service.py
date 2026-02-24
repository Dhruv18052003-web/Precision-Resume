# backend/src/services/docx_service.py

from io import BytesIO
from docx import Document
from docx.shared import Pt


def generate_docx(refined_resume_text: str) -> BytesIO:
    """
    Generate a DOCX file from refined resume text.
    Returns a BytesIO buffer containing the DOCX.
    """

    # Create Word document
    document = Document()

    # Set default font style (ATS-safe)
    style = document.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Split resume text into lines
    lines = refined_resume_text.split("\n")

    for line in lines:
        # Add empty line for spacing
        if line.strip() == "":
            document.add_paragraph("")
        else:
            para = document.add_paragraph(line)
            para_format = para.paragraph_format
            para_format.space_after = Pt(6)

    # Save document to memory
    buffer = BytesIO()
    document.save(buffer)

    # Reset buffer cursor
    buffer.seek(0)

    return buffer
