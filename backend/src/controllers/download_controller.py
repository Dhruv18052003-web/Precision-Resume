from flask import Blueprint, request, send_file
from io import BytesIO

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.units import inch

from docx import Document

download_bp = Blueprint("download", __name__)


def parse_resume_lines(resume_text: str):
    """
    Parses resume text into structured blocks.
    Very simple and ATS-safe.
    """
    lines = [line.strip() for line in resume_text.split("\n") if line.strip()]
    blocks = []

    for line in lines:
        if line.startswith("**") and line.endswith("**"):
            blocks.append(("heading", line.replace("**", "")))
        elif line.startswith("-"):
            blocks.append(("bullet", line[1:].strip()))
        else:
            blocks.append(("text", line))

    return blocks


@download_bp.route("/download-pdf", methods=["POST"])
def download_pdf():
    data = request.get_json()
    resume_text = data.get("refined_resume", "")

    if not resume_text:
        return {"error": "No resume content provided"}, 400

    buffer = BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=40,
        leftMargin=40,
        topMargin=40,
        bottomMargin=40,
    )

    styles = getSampleStyleSheet()

    # ---------- Styles ----------
    name_style = ParagraphStyle(
        "NameStyle",
        fontSize=18,
        leading=22,
        spaceAfter=10,
        fontName="Helvetica-Bold",
    )

    contact_style = ParagraphStyle(
        "ContactStyle",
        fontSize=9,
        textColor="#444444",
        spaceAfter=14,
    )

    section_style = ParagraphStyle(
        "SectionStyle",
        fontSize=12,
        fontName="Helvetica-Bold",
        spaceBefore=14,
        spaceAfter=6,
    )

    body_style = ParagraphStyle(
        "BodyStyle",
        fontSize=10,
        leading=14,
        spaceAfter=4,
    )

    bullet_style = ParagraphStyle(
        "BulletStyle",
        fontSize=10,
        leading=14,
        leftIndent=12,
    )

    # ---------- Parsing ----------
    lines = [l.strip() for l in resume_text.split("\n") if l.strip()]
    story = []

    name_done = False
    contact_buffer = []
    in_contact = False

    for i, line in enumerate(lines):
        # NAME (first bold line)
        if not name_done and line.startswith("**") and line.endswith("**"):
            name = line.replace("**", "")
            story.append(Paragraph(name, name_style))
            name_done = True
            in_contact = True
            continue

        # CONTACT INFO (collect compactly)
        if in_contact and ("@" in line or "Phone" in line or "LinkedIn" in line):
            clean = line.replace("(Optional)", "").replace("**", "")
            contact_buffer.append(clean)
            continue

        if in_contact:
            contact_line = " | ".join(contact_buffer)
            if contact_line:
                story.append(Paragraph(contact_line, contact_style))
            in_contact = False

        # SECTION HEADERS
        if line.startswith("**") and line.endswith("**"):
            section = line.replace("**", "")
            story.append(Spacer(1, 6))
            story.append(Paragraph(section.upper(), section_style))
            story.append(Spacer(1, 4))
            continue

        # BULLETS
        if line.startswith("-"):
            bullet = line[1:].strip()
            story.append(
                ListFlowable(
                    [Paragraph(bullet, bullet_style)],
                    bulletType="bullet",
                    start="circle",
                    leftIndent=14,
                )
            )
            continue

        # NORMAL TEXT
        story.append(Paragraph(line, body_style))

    doc.build(story)
    buffer.seek(0)

    return send_file(
        buffer,
        mimetype="application/pdf",
        as_attachment=True,
        download_name="refined_resume.pdf",
    )

@download_bp.route("/download-docx", methods=["POST"])
def download_docx():
    data = request.get_json()
    resume_text = data.get("refined_resume", "")

    if not resume_text:
        return {"error": "No resume content provided"}, 400

    doc = Document()
    parsed = parse_resume_lines(resume_text)

    for block_type, content in parsed:
        if block_type == "heading":
            doc.add_heading(content, level=1)

        elif block_type == "bullet":
            doc.add_paragraph(content, style="List Bullet")

        else:
            doc.add_paragraph(content)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name="refined_resume.docx",
    )
